from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "Relatorio_Testes_UI_Hugo_Marques.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_run(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.bold = True
    if level == 1:
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(46, 116, 181)
    elif level == 2:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(46, 116, 181)
    else:
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_command_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Comando", "Target", "Value", "Explicacao"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "F2F4F7")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        style_run(run, bold=True, size=10)

    for command, target, value, explanation in rows:
        cells = table.add_row().cells
        values = [command, target, value, explanation]
        for i, text in enumerate(values):
            p = cells[i].paragraphs[0]
            run = p.add_run(text)
            style_run(run, size=9)
            if i in (0, 1, 2):
                run.font.name = "Consolas"

    set_table_width(table, [1.05, 2.05, 1.25, 2.15])
    doc.add_paragraph()


def add_summary_table(doc):
    tests = [
        (
            "UI-01",
            "Abrir Login sem mudar URL",
            "Confirmar que o ecra de Login e apresentado na interface SPA.",
            "Formulario de login visivel atraves do campo de email.",
            "Passou",
        ),
        (
            "UI-02",
            "Abrir Registo sem mudar URL",
            "Confirmar que o ecra de Registo e apresentado na interface SPA.",
            "Formulario de registo visivel atraves do campo de nome.",
            "Passou",
        ),
        (
            "UI-03",
            "Login invalido mostra erro",
            "Validar tratamento de credenciais incorretas.",
            "Mensagem 'Email ou Password incorretos' apresentada.",
            "Passou",
        ),
        (
            "UI-04",
            "Registo exige campos obrigatorios",
            "Verificar que o formulario de registo nao avanca vazio.",
            "Utilizador permanece no formulario de registo.",
            "Passou",
        ),
        (
            "UI-05",
            "Alternar de Login para Registo",
            "Validar navegacao client-side entre os formularios.",
            "Formulario de registo apresentado depois de clicar Registar.",
            "Passou",
        ),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["ID", "Teste", "Objetivo", "Resultado esperado", "Estado"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "F2F4F7")
        run = cell.paragraphs[0].add_run(header)
        style_run(run, bold=True, size=9)
    for row in tests:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            run = cells[i].paragraphs[0].add_run(value)
            style_run(run, size=9)
    set_table_width(table, [0.55, 1.45, 1.75, 2.0, 0.75])
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("Engenharia de Software 2 | Hugo Simão Marques Nº33375 | IPVC ESTG")
    style_run(hr, bold=True, size=10, color="1F4D78")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Relatorio de Testes de UI")
    style_run(tr, bold=True, size=20, color="1F4D78")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Aplicacao Event Horizon - Selenium IDE")
    style_run(sr, size=12, color="555555")

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    metadata = [
        ("Disciplina", "Engenharia de Software 2"),
        ("Aluno", "Hugo Simão Marques Nº33375"),
        ("Escola", "IPVC ESTG"),
    ]
    for i, (label, value) in enumerate(metadata):
        cells = meta.rows[i].cells
        set_cell_shading(cells[0], "F2F4F7")
        style_run(cells[0].paragraphs[0].add_run(label), bold=True)
        style_run(cells[1].paragraphs[0].add_run(value))
    set_table_width(meta, [1.6, 4.9])
    doc.add_paragraph()

    add_heading(doc, "1. Introducao", 1)
    doc.add_paragraph(
        "Este documento apresenta os testes de interface realizados sobre a aplicacao Event Horizon. "
        "A parte testada corresponde aos fluxos de Login e Registo na interface principal da aplicacao."
    )
    doc.add_paragraph(
        "A aplicacao usa uma interface do tipo Single Page Application (SPA) na pagina principal, "
        "onde diferentes ecras sao apresentados por navegacao client-side, sem carregamento completo "
        "de uma nova pagina. Os testes foram executados com Selenium IDE no browser Zen."
    )

    add_heading(doc, "2. Ferramentas Utilizadas", 1)
    for item in [
        "Selenium IDE: gravacao e execucao dos testes de UI.",
        "Zen Browser: browser utilizado para executar a aplicacao e a extensao Selenium IDE.",
        "Aplicacao local: http://localhost:5007/",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    add_heading(doc, "3. Resumo dos Testes", 1)
    add_summary_table(doc)

    add_heading(doc, "4. Testes Realizados e Comandos", 1)

    add_heading(doc, "UI-01 - Abrir Login sem mudar URL", 2)
    doc.add_paragraph(
        "Objetivo: confirmar que, ao clicar em Login na pagina principal, o formulario de autenticacao "
        "e apresentado na interface dinamica da aplicacao."
    )
    add_command_table(
        doc,
        [
            ("open", "http://localhost:5007/", "", "Abre a pagina inicial da aplicacao local."),
            ("click", "xpath=//button[normalize-space()='Login']", "", "Clica no botao Login da navbar da SPA."),
            ("assert element present", 'css=input[type="email"]', "", "Confirma que existe um campo de email, indicando que o formulario de login abriu."),
        ],
    )

    add_heading(doc, "UI-02 - Abrir Registo sem mudar URL", 2)
    doc.add_paragraph(
        "Objetivo: confirmar que o formulario de registo e apresentado a partir da pagina principal, "
        "mantendo a navegacao dentro da interface dinamica."
    )
    add_command_table(
        doc,
        [
            ("open", "http://localhost:5007/", "", "Abre a pagina inicial da aplicacao."),
            ("click", "xpath=//button[normalize-space()='Registo']", "", "Clica no botao Registo."),
            ("assert element present", "xpath=//input[@placeholder='O teu nome']", "", "Confirma que existe o campo de nome do formulario de registo."),
        ],
    )

    add_heading(doc, "UI-03 - Login invalido mostra erro", 2)
    doc.add_paragraph(
        "Objetivo: validar que a aplicacao apresenta uma mensagem de erro quando o utilizador tenta "
        "autenticar-se com credenciais incorretas."
    )
    add_command_table(
        doc,
        [
            ("open", "http://localhost:5007/", "", "Abre a pagina inicial."),
            ("click", "xpath=//button[normalize-space()='Login']", "", "Mostra o formulario de login."),
            (
                "execute script",
                "const e=document.querySelector('input[type=\"email\"]'); ... dispatchEvent(new Event('input',{bubbles:true}));",
                "",
                "Preenche o campo de email por JavaScript e dispara o evento input para o React atualizar o estado.",
            ),
            (
                "execute script",
                "const p=document.querySelector('input[type=\"password\"]'); ... dispatchEvent(new Event('input',{bubbles:true}));",
                "",
                "Preenche o campo de password e dispara o evento input.",
            ),
            ("click", "xpath=//button[contains(.,'Entrar')]", "", "Submete o formulario de login."),
            ("pause", "1000", "", "Espera um segundo para a resposta da API e atualizacao da UI."),
            ("assert element present", "xpath=//*[contains(.,'Email ou Password incorretos')]", "", "Confirma que a mensagem de erro foi apresentada."),
        ],
    )

    add_heading(doc, "UI-04 - Registo exige campos obrigatorios", 2)
    doc.add_paragraph(
        "Objetivo: verificar que o formulario de registo nao e submetido quando os campos obrigatorios "
        "estao vazios."
    )
    add_command_table(
        doc,
        [
            ("open", "http://localhost:5007/", "", "Abre a aplicacao."),
            ("click", "xpath=//button[normalize-space()='Registo']", "", "Abre o formulario de registo."),
            ("click", "xpath=//button[contains(.,'Registar')]", "", "Tenta submeter o formulario sem preencher os campos."),
            ("assert element present", "xpath=//input[@placeholder='O teu nome']", "", "Confirma que o utilizador permanece no formulario de registo."),
        ],
    )

    add_heading(doc, "UI-05 - Alternar de Login para Registo", 2)
    doc.add_paragraph(
        "Objetivo: testar a navegacao client-side entre os ecras de Login e Registo."
    )
    add_command_table(
        doc,
        [
            ("open", "http://localhost:5007/", "", "Abre a pagina inicial."),
            ("click", "xpath=//button[normalize-space()='Login']", "", "Abre o formulario de login."),
            ("assert element present", 'css=input[type="email"]', "", "Confirma que o formulario de login esta visivel."),
            ("click", "xpath=//*[normalize-space()='Registar']", "", "Clica na opcao para mudar para o formulario de registo."),
            ("assert element present", "xpath=//input[@placeholder='O teu nome']", "", "Confirma que o formulario de registo ficou visivel."),
        ],
    )

    add_heading(doc, "5. Conclusao", 1)
    doc.add_paragraph(
        "Os testes executados validam funcionalidades essenciais da interface de autenticacao e registo: "
        "abertura dos formularios, navegacao dinamica, validacao de campos obrigatorios e apresentacao "
        "de mensagens de erro. Todos os testes definidos foram executados no Selenium IDE com resultado positivo."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
